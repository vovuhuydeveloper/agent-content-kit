"""
Document Parser — Extract text from PDF, DOCX, and TXT files.
Used by ContentFetcherAgent when source is an uploaded document.
"""

import logging
from pathlib import Path
from typing import Any, Dict

logger = logging.getLogger("agent.document_parser")


def parse_document(file_path: str) -> Dict[str, Any]:
    """
    Parse a document file and extract text content.

    Supports:
    - PDF (.pdf)
    - Word (.docx)
    - Text (.txt, .md)

    Returns:
        Dict with keys: title, body_text, metadata, source_type
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"Document not found: {file_path}")

    ext = path.suffix.lower()

    if ext == ".pdf":
        return _parse_pdf(path)
    elif ext == ".docx":
        return _parse_docx(path)
    elif ext in (".txt", ".md", ".text"):
        return _parse_text(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: .pdf, .docx, .txt, .md")


def _parse_pdf(path: Path) -> Dict[str, Any]:
    """Parse PDF file using PyPDF2"""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise RuntimeError("PyPDF2 not installed. Run: pip install PyPDF2")

    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text.strip())

    body_text = "\n\n".join(pages)

    # Extract metadata
    meta = reader.metadata or {}
    title = meta.get("/Title", "") or path.stem
    author = meta.get("/Author", "")

    logger.info(f"Parsed PDF: {path.name} — {len(reader.pages)} pages, {len(body_text)} chars")

    return {
        "source_url": str(path),
        "source_type": "document",
        "title": str(title),
        "description": f"PDF document: {path.name}",
        "body_text": body_text[:15000],  # Limit to avoid token overflow
        "transcript": "",
        "metadata": {
            "file_type": "pdf",
            "page_count": len(reader.pages),
            "author": str(author),
            "file_size": path.stat().st_size,
        },
    }


def _parse_docx(path: Path) -> Dict[str, Any]:
    """Parse DOCX file using python-docx"""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    doc = Document(str(path))

    # Extract paragraphs
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    body_text = "\n\n".join(paragraphs)

    # Try to get title from first heading or first paragraph
    title = path.stem
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            title = para.text.strip()
            break

    # Extract tables as text too
    table_texts = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        table_texts.append("\n".join(rows))

    if table_texts:
        body_text += "\n\n--- Tables ---\n" + "\n\n".join(table_texts)

    logger.info(f"Parsed DOCX: {path.name} — {len(paragraphs)} paragraphs, {len(body_text)} chars")

    return {
        "source_url": str(path),
        "source_type": "document",
        "title": title,
        "description": f"Word document: {path.name}",
        "body_text": body_text[:15000],
        "transcript": "",
        "metadata": {
            "file_type": "docx",
            "paragraph_count": len(paragraphs),
            "table_count": len(doc.tables),
            "file_size": path.stat().st_size,
        },
    }


def _parse_text(path: Path) -> Dict[str, Any]:
    """Parse plain text file"""
    body_text = path.read_text(encoding="utf-8", errors="ignore")

    # Use first non-empty line as title
    lines = [line.strip() for line in body_text.split("\n") if line.strip()]
    title = lines[0] if lines else path.stem

    # Strip markdown heading markers
    if title.startswith("#"):
        title = title.lstrip("#").strip()

    logger.info(f"Parsed TXT: {path.name} — {len(body_text)} chars")

    return {
        "source_url": str(path),
        "source_type": "document",
        "title": title,
        "description": f"Text document: {path.name}",
        "body_text": body_text[:15000],
        "transcript": "",
        "metadata": {
            "file_type": path.suffix.lstrip("."),
            "line_count": len(lines),
            "file_size": path.stat().st_size,
        },
    }
