"""
Shared test fixtures for agent-content-kit tests.
"""

import os
import sys
import pytest
import tempfile
import shutil
from pathlib import Path
from unittest.mock import MagicMock, patch

# Add project root to path
PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))


@pytest.fixture
def tmp_job_dir(tmp_path):
    """Create a temporary job directory"""
    job_dir = tmp_path / "test_job"
    job_dir.mkdir()
    return str(job_dir)


@pytest.fixture
def sample_context(tmp_job_dir):
    """Create a sample pipeline context"""
    return {
        "job_id": "test-job-001",
        "job_dir": tmp_job_dir,
        "source_url": "https://example.com/article",
        "language": "vi",
        "video_count": 1,
        "platforms": ["tiktok"],
        "niche": "education",
        "character_images": [],
        "competitor_urls": [],
        "agent_results": [],
        "errors": [],
        "last_checkpoint": "",
    }


@pytest.fixture
def sample_content_data():
    """Sample content data from fetcher"""
    return {
        "source_url": "https://example.com/article",
        "source_type": "webpage",
        "title": "Test Article Title",
        "description": "A test article about testing",
        "body_text": "This is the body text of the article. It contains valuable information about Python testing and best practices.",
        "transcript": "",
        "metadata": {},
    }


@pytest.fixture
def sample_scripts():
    """Sample generated scripts"""
    return [
        {
            "script_id": 1,
            "title": "Test Video Title",
            "hook": "Did you know this amazing fact?",
            "scenes": [
                {
                    "scene_id": 1,
                    "text": "Welcome! Today we learn about testing",
                    "duration": 5,
                    "visual": "A colorful intro screen",
                    "character_pose": "waving",
                },
                {
                    "scene_id": 2,
                    "text": "Testing is very important in software development",
                    "duration": 8,
                    "visual": "Code on screen",
                    "character_pose": "explaining",
                },
            ],
            "cta": "Like and subscribe for more!",
            "hashtags": ["python", "testing", "coding"],
            "estimated_duration": 45,
        }
    ]


@pytest.fixture
def mock_llm_client():
    """Mock LLM client that returns predefined responses"""
    mock = MagicMock()
    mock.generate.return_value = "Generated text response"
    mock.generate_json.return_value = {
        "script_id": 1,
        "title": "Test Video",
        "hook": "Amazing hook!",
        "scenes": [],
        "cta": "Subscribe!",
        "hashtags": ["test"],
    }
    return mock


@pytest.fixture
def sample_pdf(tmp_path):
    """Create a sample PDF file"""
    pdf_path = tmp_path / "test.pdf"
    try:
        from PyPDF2 import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        # Note: PdfWriter can't add text easily, so this creates a blank PDF
        with open(pdf_path, "wb") as f:
            writer.write(f)
    except ImportError:
        # Fallback: create a minimal PDF
        pdf_path.write_bytes(
            b"%PDF-1.4\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
            b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
            b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R>>endobj\n"
            b"xref\n0 4\n0000000000 65535 f \ntrailer<</Size 4/Root 1 0 R>>\nstartxref\n0\n%%EOF"
        )
    return str(pdf_path)


@pytest.fixture
def sample_docx(tmp_path):
    """Create a sample DOCX file"""
    docx_path = tmp_path / "test.docx"
    try:
        from docx import Document

        doc = Document()
        doc.add_heading("Test Document", level=1)
        doc.add_paragraph("This is a test paragraph with content about Python.")
        doc.add_paragraph("Another paragraph with more testing information.")
        doc.save(str(docx_path))
    except ImportError:
        # Skip if python-docx not installed
        pytest.skip("python-docx not installed")
    return str(docx_path)


@pytest.fixture
def sample_txt(tmp_path):
    """Create a sample text file"""
    txt_path = tmp_path / "test.txt"
    txt_path.write_text(
        "# Test Document\n\n"
        "This is a test text file.\n"
        "It contains information about Python programming.\n"
        "Multiple lines of valuable content.\n",
        encoding="utf-8",
    )
    return str(txt_path)


@pytest.fixture(autouse=True)
def env_setup():
    """Set up test environment variables"""
    os.environ.setdefault("CELERY_ALWAYS_EAGER", "True")
    os.environ.setdefault("DATABASE_URL", "sqlite:///./data/test.db")
    yield
